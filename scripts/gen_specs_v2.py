#!/usr/bin/env python3
"""Generate updated specs DOCX for SAO Alpha project — includes all latest features."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = '/home/z/my-project/download/SAO Alpha - Specifiche Aggiornate.docx'
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(6)

SAO_BLUE = RGBColor(0x2B, 0x73, 0xB3)

def h(text, level=1):
    return doc.add_heading(text, level=level)

def p(text, bold=False, italic=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    return para

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x3a)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hd
        for pp in cell.paragraphs:
            for r in pp.runs: r.bold = True; r.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for pp in cell.paragraphs:
                for r in pp.runs: r.font.size = Pt(9)
    doc.add_paragraph()

# === COVER ===
t = doc.add_heading('SAO Alpha', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs: r.font.color.rgb = SAO_BLUE; r.font.size = Pt(36)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('Sword Art Online — Web RPG'); r.italic = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x5C,0xC4,0xF0)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('Specifiche Tecniche Complete — Versione Aggiornata'); r.font.size = Pt(13)
doc.add_paragraph()
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('Versione: Alpha 1.0 (Aggiornata)\nData: 26 Giugno 2026\nRepository: github.com/emanuelpapasidero98-crypto/Project-SAO-Alpha'); r.font.size = Pt(11)
doc.add_page_break()

# === 1. PANORAMICA ===
h('1. Panoramica del Progetto')
p('SAO Alpha è una web app RPG a turni ispirata a Sword Art Online. Il progetto riproduce fedelmente l\'interfaccia UI canonica dell\'anime usando ESCLUSIVAMENTE asset grafici e audio dai due repository GitHub dell\'utente.')
p('Regola d\'oro: NON CREARE MAI GRAFICHE DA ZERO. Usa SOLO asset dai repo:', bold=True)
code('Repo 1: https://github.com/emanuelpapasidero98-crypto/asset-gioco-di-SAO.git\nRepo 2: https://github.com/emanuelpapasidero98-crypto/Progetto-SAO.git')
p('Stato attuale:', bold=True)
bullet('Login screen con boot sequence VR NerveGear')
bullet('Transizione "Link Start" cinematografica')
bullet('Creazione personaggio (genere + nome + riepilogo statistiche)')
bullet('Splash screen "Entering Aincrad"')
bullet('Game screen con HUD completo (barre HP/MP/Energia + LV)')
bullet('Menu a cascata SAO (9 voci con icone dedicate)')
bullet('Finestra notifiche canonica SAO (3 strati SVG + bottoni blu/rosso)')
bullet('Character Panel con VR hover + box equipaggiamento (5 slot)')
bullet('Inventory Panel (12 categorie, ricerca, ordinamento, equip/move)')
bullet('Bag Panel (max 10 slot, equip/move)')
bullet('ItemDetailModal (icona rotante VR + descrizione)')
bullet('Floor Panel (selezione Piano 1 → zone Città/Plains)')
bullet('Stats engine completo (milestone, XP table, derived stats)')
bullet('Background particellare 3D interattivo (React Three Fiber)')
doc.add_page_break()

# === 2. STACK ===
h('2. Stack Tecnologico')
table(['Categoria','Tecnologia','Versione'], [
    ['Framework','Next.js','16.1.3 (App Router, Turbopack)'],
    ['Linguaggio','TypeScript','5.x (strict)'],
    ['Styling','Tailwind CSS','4.x + shadcn/ui'],
    ['3D','React Three Fiber','9.6.1 + drei + three@0.184'],
    ['Animazioni','Framer Motion','12.23.2'],
    ['Font','SAO UI (SAOUI-Regular.otf)','Dal repo Progetto-SAO/Fonts/'],
    ['Package Manager','Bun','1.3.14'],
])
doc.add_page_break()

# === 3. FLUSSO ===
h('3. Flusso di Gioco (5 Stage)')
code('login → linkstart → creation → entering → game\n                                         ↓ (Log Out)\n                                       login')
p('Stage 1 - Login: boot VR multi-stage, NerveGear gigante, Tech-Circle rotante, pannello SAO_Login_UI_Transparent.svg, input testo scuro, pulsante LINK START')
p('Stage 2 - Link Start: flash full-screen con anelli concentrici, audio LinkStart.SAO.Kirito.wav, durata 2.6s')
p('Stage 3 - Character Creation: scelta genere (SAO_Man/Woman.svg) → riepilogo con nome + stats (tutte a 1). SAO è skill-based, NESSUNA CLASSE.')
p('Stage 4 - Entering Aincrad: NerveGear gigante, "SWORD ART ONLINE VRMMORPG" + badge "V ALPHA 1.0", "BENVENUTO AD AINCRAD", durata 4.5s')
p('Stage 5 - Game: background Aincrad.png, HUD top-left, menu top-right, notification window, crosshair VR')
doc.add_page_break()

# === 4. STRUTTURA ===
h('4. Struttura del Codice')
code('''src/
├── app/
│   ├── layout.tsx          # Font SAO UI globale, metadata
│   ├── page.tsx            # State machine (5 stage)
│   └── globals.css         # @font-face + scrollbar SAO
├── components/sao/
│   ├── LoginScreen.tsx
│   ├── CharacterCreation.tsx
│   ├── GameScreen.tsx      # Orchestrazione game UI
│   ├── SaoHUD.tsx          # Barre HP/MP/Energia (clip-path LV box)
│   ├── SaoMainMenu.tsx     # Menu cascata 9 voci con icone
│   ├── SaoNotificationWindow.tsx # Finestra notifiche canonica
│   ├── CharacterPanel.tsx  # Scheda personaggio + equip
│   ├── SaoPanel.tsx        # Borsa/Inventario condiviso
│   ├── ItemDetailModal.tsx # Ispezione oggetto VR (rotante)
│   ├── FloorPanel.tsx      # Selezione piano + zone
│   └── ParticleBackground.tsx
├── hooks/
│   └── useSaoSound.ts      # 18 suoni SAO
└── lib/
    ├── sao-data.ts          # Re-export + gender data
    ├── sao-stats.ts         # Stats engine completo
    ├── sao-stats-engine.ts  # Milestone, XP, derived stats
    ├── sao-types.ts         # PlayerStats, StatBonus, Player
    ├── sao-inventory-types.ts # Item, EquipmentState, CATEGORIES
    └── sao-sample-items.ts  # 12 oggetti di esempio''')
doc.add_page_break()

# === 5. COMPONENTI ===
h('5. Componenti SAO — Dettagli')

h('5.1 SaoHUD (Barre HP/MP/Energia)', level=2)
bullet('Template: [Blank] 2.png (1620x258) con forma esagonale')
bullet('3 versioni: blank-hp.png (verde), blank-mp.png (blu), blank-energy.png (giallo)')
bullet('Valori current/max DENTRO i box semi-trasparenti esistenti del PNG (current a X=66.3%, max a X=78.3%, top 76.4%)')
bullet('LV SOLO sulla barra Energy (X=93.2%) — numero intero, colore bianco')
bullet('HP/MP: clip-path nasconde il box LV del PNG (polygon 0% 0%, 100% 0%, 100% 52%, 84% 52%, 84% 100%, 0% 100%)')
bullet('Nome giocatore: solo testo (no box), bianco con glow cyan, sopra la barra HP')
bullet('Nessun suono hover sulle barre')
bullet('Valori come numeri interi (120, non 0120)')

h('5.2 SaoMainMenu (Menu cascata)', level=2)
bullet('Singola icona Config.svg in alto a destra → apre cascata 9 voci')
bullet('9 voci: Personaggio, Borsa, Inventario, Quest, Piano, Party, Opzioni, Messaggi, Log Out')
bullet('Ogni voce ha icona dedicata (Player.svg, Items.svg, Equipment.svg, ecc.) con variante _on per hover')
bullet('Animazione: opacity 0→1 (400ms) + y -height→0 (600ms OutQuart) + suono popupMenu (hover voci)')
bullet('Suono apertura = suono chiusura (dismissLauncher)')
bullet('Nessun suono hover sulle CARD GRANDI (Personaggio/Borsa/Inventario/Floor)')

h('5.3 SaoNotificationWindow (Finestra notifiche)', level=2)
bullet('Usa 3 SVG canonici: Pezzo superiore (header), Parte interna (body), Parete sotto (footer con bottoni)')
bullet('Bottoni blu/rosso: NESSUNA animazione hover, NESSUNA scritta (si capiscono da soli)')
bullet('Messaggi NON SCADONO — rimangono fino a chiusura manuale')
bullet('Animazione: opacity 0→1 (400ms) + y -height→0 (600ms OutQuart)')

h('5.4 CharacterPanel (Scheda personaggio)', level=2)
bullet('Card bianca SAO con corner tagliati + hover VR (3D tilt 8deg + glow)')
bullet('Sinistra: avatar in box scuro SAO + nome + LV (nero grassetto rilievo) + sub-stats + box EQUIPAGGIAMENTO (5 slot)')
bullet('Destra: XP bar (numerica, no percentuale) + 7 stats con icone + derived stats + resistenze + bonus utility')
bullet('5 slot equip: ARMA, SCUDO, ARMATURA, ACC.1, ACC.2 (item equipaggiati cliccabili → ItemDetailModal)')
bullet('Arma 2 mani → slot SCUDO BLOCCATO (rosso, opacità 0.4)')
bullet('Nessun suono hover sulla card')

h('5.5 SaoPanel (Borsa + Inventario)', level=2)
bullet('Componente condiviso per Borsa (max 10 slot) e Inventario (tutti gli oggetti)')
bullet('Inventario: 12 categorie tab + ricerca + ordinamento (recente/A-Z)')
bullet('Card oggetti con hover VR (3D tilt 12deg + scale 1.03)')
bullet('Pulsanti EQUIP e BORSA/INV. su ogni oggetto')
bullet('Nessuna notifica di conferma equip/disequip/sposta')
bullet('Nomi oggetti in rilievo (fontWeight 700 + text-shadow)')

h('5.6 ItemDetailModal (Ispezione VR)', level=2)
bullet('Click icona oggetto → modal con icona INGRANDITA che ROTA (8s/giro)')
bullet('Card con hover VR (3D tilt + glow), stessa estetica delle altre card')
bullet('Mostra: categoria, nome (grassetto rilievo), descrizione, bonus stats, equipaggiabilità')
bullet('Nessun suono hover sulla card')

h('5.7 FloorPanel (Selezione piano)', level=2)
bullet('Card grande (min 1100px) con molto respiro')
bullet('Vista 1: immagine grande Aincrad.png cliccabile per Piano 1 + piani 2-4 bloccati')
bullet('Vista 2: 2 zone (Città degli Inizi con Città dell\'inizio.png + Pianure dell\'inizio con Pianure dell\'inizio.svg)')
bullet('Badge CITTÀ/ESPLORA su ogni zona')
bullet('Hover VR sulla card (3D tilt + glow)')

h('5.8 ParticleBackground', level=2)
bullet('1500 particelle con shader custom (glow + twinkle)')
bullet('2 point lights che seguono il mouse (cyan + bianco)')
bullet('Repulsione particelle vicino al cursore')
doc.add_page_break()

# === 6. STATS ENGINE ===
h('6. Stats Engine Completo')
p('Sistema statistiche canonico SAO in src/lib/sao-stats.ts:')
bullet('7 statistiche: STR, DEX, AGI, VIT, RES, MEN, INT')
bullet('10 milestone per stat (5,10,20,30,40,50,60,70,80,90 punti) con bonus cumulativi')
bullet('XP table fissa livelli 1-90 (100 XP al lv1 → 400500 XP al lv89)')
bullet('HP: base 120, +5%/livello composto + bonus VIT')
bullet('MP: base 50, +5%/livello + bonus MEN')
bullet('SP: base 60, +5%/livello + bonus RES')
bullet('Punti stat per livello: 3 (lv2-10), 2 (lv11-40), 1 (lv41-90), 0 (cap 90)')
bullet('Derived stats: attack, defense, dodge, crit, stun, resistenze status, XP/Col multiplier')
bullet('Statistiche iniziali: tutte a 1 (SAO è skill-based, niente classi)')
doc.add_page_break()

# === 7. INVENTORY ===
h('7. Sistema Inventario + Borsa + Equipaggiamento')
table(['Aspetto','Dettaglio'], [
    ['Categorie','12: Spade 1h, Spadoni, Asce 1h, Asce 2h, Pugnali, Fioretto, Scudi, Armature, Accessori, Oggetti, Pozioni, Oggetti Missione'],
    ['Borsa','Max 10 oggetti trasportabili (per esplorazione futura)'],
    ['Equipaggiamento','5 slot: ARMA, SCUDO, ARMATURA, ACC.1, ACC.2'],
    ['Regola 2 mani','Arma a 2 mani blocca lo slot SCUDO'],
    ['Spostamento','Inventario <-> Borsa (solo se c\'e spazio)'],
    ['Ricerca','Barra di testo per nome oggetto'],
    ['Ordinamento','Recente (acquiredAt) o Alfabetico (A-Z)'],
    ['ItemDetailModal','Click icona → modal VR con icona rotante + descrizione'],
    ['Notifiche','NESSUNA notifica di conferma equip/disequip/sposta'],
])
doc.add_page_break()

# === 8. ASSET ===
h('8. Asset Utilizzati (SOLO dai repo GitHub)')
p('TUTTI gli asset provengono da:')
code('Repo 1: https://github.com/emanuelpapasidero98-crypto/asset-gioco-di-SAO.git\nRepo 2: https://github.com/emanuelpapasidero98-crypto/Progetto-SAO.git')
p('Asset principali:', bold=True)
bullet('Font: SAOUI-Regular.otf (Progetto-SAO/Fonts/)')
bullet('Login: SAO_Nervegear.svg, SAO_Login_UI_Transparent.svg, SAO_Login_Tech-Circle_v1.000.svg')
bullet('Avatar: SAO_Man.svg, SAO_Woman.svg (NON Avatar_maschio/femmina.svg)')
bullet('Barre: [Blank] 2.png + contenitore/Contenuto/Contorno/pezzi valori SVG')
bullet('Finestra: Pezzo superiore + Parte interna + Parete sotto + Pulsante azzurro/rosso SVG')
bullet('Menu icone: Player, Items, Equipment, Quest & Message Box, Location, Party & Profile, Option, Message, Logout (con _on varianti)')
bullet('Stats: 7 PNG (Forza, Vita, Agilità, Destrezza, Intelligenza, Mente, Resistenza)')
bullet('Equipment: 24 PNG (icon_sword, icon_bow, icon_dagger, ecc.)')
bullet('Backgrounds: Aincrad.png, Città dell\'inizio.png, Pianure dell\'inizio.svg')
bullet('Audio: 32 WAV (Startup, LinkStart, Click, Popup, Notify, Dismiss, ecc.)')
doc.add_page_break()

# === 9. ANIMAZIONI ===
h('9. Animazioni e Interazioni')
p('Animazioni canoniche SAO (da MenuView.qml / PanelView.qml):', bold=True)
code('OPEN: opacity 0→1 (400ms) + y -height→0 (600ms OutQuart) + sound after 300ms\nCLOSE: opacity 1→0 (400ms) + y 0→-height (300ms InQuad)')
p('Hover VR su card grandi (Personaggio/Borsa/Inventario/Floor/ItemDetail):', bold=True)
bullet('3D tilt rotateX/Y (6-8deg) seguendo il mouse')
bullet('Glow radiale cyan che segue il cursore')
bullet('Box-shadow che si intensifica')
bullet('NESSUN suono hover sulle card grandi')
p('Hover VR su card oggetti (Inventory/Bag):', bold=True)
bullet('3D tilt rotateX/Y (12deg) + scale 1.03')
p('Hover voci menu cascata:', bold=True)
bullet('Suono popupMenu su ogni voce')
bullet('Icona normale → icona _on (crossfade 200ms)')
doc.add_page_break()

# === 10. REGOLE ===
h('10. Regole Fondamentali (NON rompere MAI)')
rules = [
    'NON CREARE MAI GRAFICHE DA ZERO — usa SOLO asset dai repo GitHub',
    'Usa SEMPRE il font SAO UI (SAOUI-Regular.otf)',
    'Colori canonici SAO: #2B73B3, #FBFBFB, #EBA601, #BE2156, #7FC522',
    'Barre HP/MP/Energia: PNG [Blank] 2.png, valori DENTRO i box esistenti',
    'LV SOLO sulla barra Energy (clip-path nasconde box LV su HP/MP)',
    'SAO NON ha classi — skill-based, niente classi',
    'Animazioni menu/finestre: pattern MenuView.qml (fade + slide dall\'alto)',
    'Finestra notifica: 3 SVG canonici, NON ricrearla',
    'Bottoni blu/rosso: nessuna animazione hover, nessuna scritta',
    'Messaggi NON scadono — rimangono fino a chiusura manuale',
    'Nessuna notifica di conferma equip/disequip/sposta',
    'Nessun suono hover sulle CARD GRANDI (Personaggio/Borsa/Inventario/Floor)',
    'Suono hover SUI sotto-menu del menu cascata (popupMenu)',
    'Menu: una sola icona → cascata 9 voci con icone dedicate',
    'Font weight 400 (non grassetto), tranne nomi oggetti (700 rilievo) e LV (700 rilievo)',
    'Card SAO: corner tagliati (clip-path), NON arrotondati',
    'Scrollbar stile SAO (.sao-scroll: thin, cyan glow)',
]
for r in rules:
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run('⚠ ' + r)
    run.font.color.rgb = RGBColor(0xBE, 0x21, 0x56)
    run.bold = True
doc.add_page_break()

# === 11. ROADMAP ===
h('11. Roadmap — Prossimi Sviluppi')
table(['Priorità','Feature','Descrizione'], [
    ['ALTA','Modalità esplorazione','Click su Pianure dell\'inizio → scena esplorazione book-style'],
    ['ALTA','Combat system a turni','Sistema RPG a turni con HP/MP/Energy'],
    ['ALTA','Status effects','27 Status Icons per effetti (veleno, stun, burn, ecc.)'],
    ['MEDIA','Sistema skill','Skill tree sbloccabile (SAO è skill-based)'],
    ['MEDIA','Party system','Formazione party con NPC'],
    ['MEDIA','Quest system','Log quest con tracker'],
    ['MEDIA','Shop/NPC','Acquisti e interazioni NPC nella città'],
    ['BASSA','Salvataggio','Persistenza con Prisma/SQLite'],
])
doc.add_page_break()

# === 12. ISTRUZIONI CLAUDE ===
h('12. Istruzioni per Claude (AI assistant)')
p('Context:', bold=True)
bullet('SAO Alpha è una web app RPG che riproduce l\'UI di Sword Art Online')
bullet('Stato: login + creation + game HUD + character panel + inventory + bag + floor panel completi')
bullet('Prossimo step: modalità esplorazione o combat system')
p('Regole d\'oro:', bold=True)
bullet('NON CREARE MAI GRAFICHE DA ZERO — usa solo asset dai repo GitHub')
bullet('Le animazioni seguono il pattern MenuView.qml (fade + slide dall\'alto)')
bullet('Font SAO UI ovunque, weight 400 (non grassetto eccessivo)')
bullet('Card con corner tagliati (clip-path), non arrotondati')
bullet('Testi in italiano')
p('Asset repo URLs:', bold=True)
code('Repo 1: https://github.com/emanuelpapasidero98-crypto/asset-gioco-di-SAO.git\nRepo 2: https://github.com/emanuelpapasidero98-crypto/Progetto-SAO.git\nProject: https://github.com/emanuelpapasidero98-crypto/Project-SAO-Alpha.git')
p('File chiave da leggere prima di modificare:', bold=True)
code('''src/app/page.tsx                    # State machine
src/components/sao/GameScreen.tsx   # Schermata gioco
src/components/sao/SaoHUD.tsx       # Barre HP/MP/Energia
src/components/sao/SaoMainMenu.tsx  # Menu cascata
src/components/sao/CharacterPanel.tsx # Scheda personaggio
src/components/sao/SaoPanel.tsx     # Borsa/Inventario
src/components/sao/FloorPanel.tsx   # Selezione piano
src/lib/sao-stats.ts                # Stats engine
src/lib/sao-inventory-types.ts      # Tipi inventario''')
p('Cose da NON fare:', bold=True)
bullet('NON reintroitare classi (SAO è skill-based)')
bullet('NON usare Avatar_maschio/femmina.svg (usare SAO_Man/Woman.svg)')
bullet('NON aggiungere LV su HP/MP (solo su Energy)')
bullet('NON creare box valori per le barre (usare quelli del PNG)')
bullet('NON mettere suoni hover sulle card grandi')
bullet('NON far scadere i messaggi')
bullet('NON mettere notifiche di conferma equip/disequip/sposta')

doc.save(OUTPUT)
import os
print(f'Generato: {OUTPUT}')
print(f'Dimensione: {os.path.getsize(OUTPUT)/1024:.1f} KB')
