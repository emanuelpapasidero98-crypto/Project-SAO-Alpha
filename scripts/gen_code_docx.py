#!/usr/bin/env python3
"""
Generate a single DOCX file containing ALL the SAO Alpha source code.
This is a lightweight backup solution that works around the download size limit.
The DOCX will contain the full source code of every file, organized by directory.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = '/home/z/my-project'
OUTPUT = '/home/z/my-project/download/SAO Alpha - Codice Sorgente.docx'

# Directories and files to include
INCLUDE_DIRS = ['src', 'scripts']
INCLUDE_FILES = [
    'package.json',
    'tsconfig.json',
    'next.config.ts',
    'tailwind.config.ts',
    'postcss.config.mjs',
    'components.json',
    'eslint.config.mjs',
    'Caddyfile',
    '.env',
    'worklog.md',
]
EXCLUDE_PATTERNS = ['node_modules', '.next', '.git', '__pycache__']

def should_exclude(path):
    return any(p in path for p in EXCLUDE_PATTERNS)

def collect_files():
    files = []
    # Collect from include dirs
    for d in INCLUDE_DIRS:
        full_dir = os.path.join(PROJECT_ROOT, d)
        if not os.path.exists(full_dir):
            continue
        for root, dirs, filenames in os.walk(full_dir):
            if should_exclude(root):
                continue
            for fn in filenames:
                full_path = os.path.join(root, fn)
                rel_path = os.path.relpath(full_path, PROJECT_ROOT)
                if not should_exclude(rel_path):
                    files.append((rel_path, full_path))
    # Collect individual files
    for fn in INCLUDE_FILES:
        full_path = os.path.join(PROJECT_ROOT, fn)
        if os.path.exists(full_path):
            files.append((fn, full_path))
    # Sort by relative path
    files.sort(key=lambda x: x[0])
    return files

def read_file_content(path):
    """Read file content as text. Skip binary files."""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(path, 'r', encoding='latin-1') as f:
                return f.read()
        except:
            return f'[Binary file - {os.path.getsize(path)} bytes]'

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(8)

    # Cover/title
    title = doc.add_heading('SAO Alpha - Codice Sorgente', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Backup completo del codice sorgente del progetto SAO Alpha\n'
                    'Sword Art Online - Web RPG con UI canonica SAO')
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Summary
    files = collect_files()
    summary = doc.add_paragraph()
    summary.add_run(f'Totale file inclusi: {len(files)}\n').bold = True

    # List of files
    doc.add_heading('Indice dei file', level=1)
    for rel_path, _ in files:
        p = doc.add_paragraph(rel_path, style='List Bullet')
        p.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # Each file
    doc.add_heading('Codice sorgente', level=1)

    for rel_path, full_path in files:
        # File header
        h = doc.add_heading(rel_path, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2B, 0x73, 0xB3)  # SAO blue

        # File content
        content = read_file_content(full_path)
        # Limit very long files to first 5000 lines to keep docx manageable
        lines = content.split('\n')
        if len(lines) > 5000:
            content = '\n'.join(lines[:5000]) + f'\n\n... [truncated, {len(lines)-5000} more lines] ...'

        # Add content as monospace paragraph
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.font.name = 'Consolas'
        run.font.size = Pt(7)
        # Light gray background-ish via color
        run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x3a)

        # Spacer
        doc.add_paragraph()

    # Save
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f'Generated: {OUTPUT}')
    print(f'Size: {size/1024:.1f} KB ({size/1024/1024:.2f} MB)')
    print(f'Files included: {len(files)}')

if __name__ == '__main__':
    main()
