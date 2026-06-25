#!/usr/bin/env python3
"""
Generate a comprehensive DOCX specification document for the SAO Alpha project.
This document contains ALL project details: architecture, features, components,
assets, animations, decisions, and roadmap — suitable for sharing with another
AI assistant (Claude) to continue development.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = '/home/z/my-project/download/SAO Alpha - Specifiche Progetto.docx'

doc = Document()

# === STYLES ===
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(6)

# Helper functions
def add_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x3a)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table

SAO_BLUE = RGBColor(0x2B, 0x73, 0xB3)
SAO_CYAN = RGBColor(0x5C, 0xC4, 0xF0)

# ===================================================================
# COVER
# ===================================================================
title = doc.add_heading('SAO Alpha', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = SAO_BLUE
    run.font.size = Pt(36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sword Art Online — Web RPG')
run.italic = True
run.font.size = Pt(16)
run.font.color.rgb = SAO_CYAN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Specifiche Tecniche Complete del Progetto')
run.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Versione: Alpha 1.0\nData: 25 Giugno 2026\nRepository: github.com/emanuelpapasidero98-crypto/Project-SAO-Alpha')
run.font.size = Pt(11)

doc.add_page_break()

# ===================================================================
# INDICE
# ===================================================================
add_heading('Indice', level=1)
toc_items = [
    '1. Panoramica del Progetto',
    '2. Stack Tecnologico',
    '3. Flusso di Gioco (Stage Machine)',
    '4. Struttura del Codice',
    '5. Componenti SAO — Dettagli',
    '6. Asset Utilizzati (esclusivamente dai repo GitHub)',
    '7. Sistema Audio',
    '8. Animazioni e Interazioni',
    '9. Font SAO UI',
    '10. Palette Colori SAO',
    '11. Decisioni di Design Prese',
    '12. Regole Fondamentali (NON rompere)',
    '13. Problemi Risolti e Soluzioni',
    '14. Roadmap — Prossimi Sviluppi',
    '15. Come Clonare e Avviare il Progetto',
    '16. Istruzioni per Claude (AI assistant)',
]
for item in toc_items:
    add_bullet(item)

doc.add_page_break()

# ===================================================================
# 1. PANORAMICA
# ===================================================================
add_heading('1. Panoramica del Progetto', level=1)

add_para('SAO Alpha è una web app RPG a turni ispirata a Sword Art Online, '
         'che riproduce fedelmente l\'interfaccia UI canonica dell\'anime. '
         'Il progetto è costruito interamente con asset grafici e audio '
         'provenienti da due repository GitHub ufficiali dell\'utente, '
         'senza creare MAI grafiche da zero.')

add_para('Regola d\'oro fondamentale:', bold=True)
add_para('NON DEVI MAI CREARE GRAFICHE DA ZERO. Usa SOLO ed ESCLUSIVAMENTE '
         'gli asset (SVG, PNG, audio WAV, font) trovati nei due repository:',
         italic=True)

add_code('Repo 1 (asset grafici + audio): https://github.com/emanuelpapasidero98-crypto/asset-gioco-di-SAO.git\n'
         'Repo 2 (audio + QML templates): https://github.com/emanuelpapasidero98-crypto/Progetto-SAO.git')

add_para('Stato attuale del progetto:', bold=True)
add_bullet('Login screen con boot sequence VR NerveGear')
add_bullet('Transizione "Link Start" cinematografica')
add_bullet('Creazione personaggio (genere + nome + riepilogo statistiche)')
add_bullet('Splash screen "Entering Aincrad"')
add_bullet('Game screen con HUD completo (barre HP/MP/Energia + LV)')
add_bullet('Menu a cascata SAO (9 voci: Personaggio, Borsa, Inventario, Quest, Piano, Party, Opzioni, Messaggi, Log Out)')
add_bullet('Finestra notifiche canonica SAO con bottoni blu/rossi')
add_bullet('Background particellare 3D interattivo (React Three Fiber)')

doc.add_page_break()

# ===================================================================
# 2. STACK TECNOLOGICO
# ===================================================================
add_heading('2. Stack Tecnologico', level=1)

add_table(
    ['Categoria', 'Tecnologia', 'Versione', 'Note'],
    [
        ['Framework', 'Next.js', '16.1.3', 'App Router, Turbopack'],
        ['Linguaggio', 'TypeScript', '5.x', 'Strict mode'],
        ['Styling', 'Tailwind CSS', '4.x', 'Con shadcn/ui'],
        ['3D', 'React Three Fiber', '9.6.1', '@react-three/fiber + drei + three@0.184'],
        ['Animazioni', 'Framer Motion', '12.23.2', 'Tutte le transizioni e hover effects'],
        ['Database', 'Prisma ORM', '6.11.1', 'SQLite (non ancora usato attivamente)'],
        ['State', 'Zustand + TanStack Query', '5.x / 5.82', 'Disponibili ma non ancora usati'],
        ['Auth', 'NextAuth.js', '4.24.11', 'Disponibile ma non ancora usato'],
        ['Font', 'SAO UI (SAOUI-Regular.otf)', '—', 'Dal repo Progetto-SAO/Fonts/'],
        ['Package Manager', 'Bun', '1.3.14', 'Per install e dev server'],
        ['Porta dev server', '3000', '—', 'Auto-avviato dal sistema'],
    ]
)

doc.add_page_break()

# ===================================================================
# 3. FLOW DI GIOCO
# ===================================================================
add_heading('3. Flusso di Gioco (Stage Machine)', level=1)

add_para('Il progetto usa una state machine in src/app/page.tsx con 5 stage:')

add_code('''type Stage = 'login' | 'linkstart' | 'creation' | 'entering' | 'game';

// Flow:
login → linkstart → creation → entering → game
                                          ↓ (Log Out)
                                        login''')

add_heading('Stage 1: Login (LoginScreen.tsx)', level=2)
add_bullet('Boot sequence VR multi-stage (NerveGear startup → system → panel → welcome)')
add_bullet('Logo NerveGear gigante (SAO_Nervegear.svg) con glow pulsante')
add_bullet('Tech-Circle rotante (SAO_Login_Tech-Circle_v1.000.svg)')
add_bullet('Pannello login (SAO_Login_UI_Transparent.svg) con campi input sovrapposti')
add_bullet('Input username/password con testo SCURO (#1a2a3a) su sfondo bianco del SVG')
add_bullet('Pulsante "LINK START" con glow cyan')
add_bullet('Parallax head-tracking (NerveGear si muove col mouse)')
add_bullet('Scanlines + chromatic aberration overlay (effetto VR)')
add_bullet('VR corner markers ai 4 angoli del viewport')
add_bullet('Boot log lines che si stampano riga per riga')

add_heading('Stage 2: Link Start (LinkStartTransition)', level=2)
add_bullet('Full-screen flash con anelli concentrici espandenti')
add_bullet('Testo "LINK START" pulsante')
add_bullet('Audio: LinkStart.SAO.Kirito.wav')
add_bullet('Durata: 2.6 secondi')

add_heading('Stage 3: Character Creation (CharacterCreation.tsx)', level=2)
add_bullet('Step 1: Scelta genere (UOMO/DONNA) — card SAO spigolose con VR hover')
add_bullet('Step 2: Riepilogo con input nome + statistiche (tutte a 1)')
add_bullet('Avatar: SAO_Man.svg / SAO_Woman.svg (NON usare Avatar_maschio.svg/femmina.svg)')
add_bullet('Card con clip-path angolare (corner tagliati a 16px)')
add_bullet('3D tilt on hover + spotlight che segue il cursore')
add_bullet('Statistiche iniziali: tutte a 1 (non 8) — SAO è skill-based')

add_heading('Stage 4: Entering Aincrad (EnteringAincrad)', level=2)
add_bullet('NerveGear gigante (preponderante) con glow pulsante')
add_bullet('Anelli pulsanti concentrici attorno al NerveGear')
add_bullet('Testo "SWORD ART ONLINE VRMMORPG" + badge "V ALPHA 1.0"')
add_bullet('Testo "BENVENUTO AD AINCRAD"')
add_bullet('NON c\'è più la scritta "NERVEGEAR" bianca gigante (rimossa)')
add_bullet('NON c\'è più "il mondo di spada e magia ti attende" (rimossa)')
add_bullet('Durata: 4.5 secondi → poi va a game')

add_heading('Stage 5: Game (GameScreen.tsx)', level=2)
add_bullet('Background: Aincrad.png con overlay scuro per leggibilità')
add_bullet('HUD top-left: SaoHUD (HP/MP/Energia + LV)')
add_bullet('Menu top-right: SaoMainMenu (singola icona → cascata)')
add_bullet('Notification window: SaoNotificationWindow (modale canonica)')
add_bullet('Welcome splash iniziale ("FLOOR 1 — AINCRAD")')
add_bullet('Crosshair VR bottom-center')
add_bullet('Pulsante ESCI top-center')

doc.add_page_break()

# ===================================================================
# 4. STRUTTURA DEL CODICE
# ===================================================================
add_heading('4. Struttura del Codice', level=1)

add_code('''src/
├── app/
│   ├── layout.tsx          # Root layout: font SAO UI globale, metadata
│   ├── page.tsx            # State machine principale (5 stage)
│   └── globals.css         # @font-face SAO UI + tema Tailwind
├── components/
│   ├── sao/
│   │   ├── LoginScreen.tsx          # Schermata login con boot VR
│   │   ├── CharacterCreation.tsx    # Creazione personaggio (2 step)
│   │   ├── GameScreen.tsx           # Schermata di gioco
│   │   ├── SaoHUD.tsx              # Barre HP/MP/Energia + LV
│   │   ├── SaoMainMenu.tsx         # Menu cascata 9 voci
│   │   ├── SaoNotificationWindow.tsx # Finestra notifiche canonica
│   │   └── ParticleBackground.tsx  # Background 3D R3F
│   └── ui/                          # shadcn/ui components (pre-esistenti)
├── hooks/
│   └── useSaoSound.ts               # Audio manager (18 suoni SAO)
└── lib/
    └── sao-data.ts                  # Classi/stat personaggio (rimosse classi)''')

add_heading('File importanti in public/sao/', level=2)
add_code('''public/sao/
├── fonts/SAOUI-Regular.otf         # Font canonico SAO
├── login/
│   ├── SAO_Nervegear.svg           # Logo NerveGear
│   ├── SAO_Login_UI_Transparent.svg # Pannello login
│   └── SAO_Login_Tech-Circle_v1.000.svg # Cerchio tecnologico rotante
├── characters/
│   ├── SAO_Man.svg                 # Avatar maschile (USA QUESTO)
│   ├── SAO_Woman.svg               # Avatar femminile (USA QUESTO)
│   ├── SAO_Man.png
│   ├── Avatar_maschio.svg          # NON USARE (troppo grande, non piaciuto)
│   └── Avatar_femmina.svg          # NON USARE
├── hpbar/
│   ├── [Blank] 2.png              # Template barra SAO (1620x258)
│   ├── blank-hp.png               # Versione verde (originale)
│   ├── blank-mp.png               # Versione blu (hue rotation)
│   ├── blank-energy.png           # Versione gialla (hue rotation)
│   ├── contenitore Barre.svg      # Sfondo scuro barra
│   ├── Contenuto barra HP.svg     # Fill verde
│   ├── Contorno barre.svg         # Cornice
│   ├── pezzi valori barre e lv.svg # Box valori con "/" e "LV:"
│   ├── values-only.svg            # Box valori senza LV (per HP/MP)
│   └── values-with-lv.svg         # Box valori con LV (per Energy)
├── window/                         # 5 SVG finestra notifiche SAO
│   ├── SAO_UI-Window_blank.svg    # Finestra completa canonica
│   ├── Pezzo superiore finestra.svg # Header gradiente
│   ├── Parte interna finestra.svg   # Body grigio
│   ├── Parete sotto della finestra.svg # Footer con bottoni
│   ├── Pulsante azzurro.svg       # Bottone OK (cerchio blu)
│   └── Pulsante rosso.SVG         # Bottone X (cerchio rosso)
├── menu/                          # 4 icone menu (off/on)
│   ├── Man.svg / Man_on.svg       # Personaggio
│   ├── Location.svg / Location_on.svg # Mappa
│   ├── Message.svg / Message_on.svg   # Messaggi
│   └── Config.svg / Config_on.svg     # Impostazioni
├── backgrounds/
│   ├── Aincrad.png                # Città volante (usato nel game)
│   ├── Città dell'inizio.png
│   └── Background Ui.png
├── stats/                         # 7 icone statistiche (4-5MB cadauna)
├── equipment/                     # 24 icone equipaggiamento
├── hex/                           # 7 SVG esagonali (Warning, You are dead, ecc.)
├── bars/                          # Pezzi barra HP/Mana/Energia
└── audio/                         # 32 file WAV audio SAO''')

doc.add_page_break()

# ===================================================================
# 5. COMPONENTI SAO - DETTAGLI
# ===================================================================
add_heading('5. Componenti SAO — Dettagli', level=1)

add_heading('5.1 SaoHUD (barre HP/MP/Energia)', level=2)
add_para('Posizione: top-left (fixed). Contiene 3 barre sovrapposte verticalmente.')
add_para('Asset usati:', bold=True)
add_bullet('[Blank] 2.png (1620x258) — template barra canonica con bordo scuro #373737, linea bianca decorativa, forma esagonale (lato destro tagliato), fill con gradiente')
add_bullet('blank-hp.png (verde #7FC522) — originale')
add_bullet('blank-mp.png (blu #2B73B3) — generato via hue rotation')
add_bullet('blank-energy.png (giallo #EBA601) — generato via hue rotation')
add_para('Slot valori integrato nel PNG:', bold=True)
add_para('Il PNG [Blank] 2.png contiene GIA\' 2 box semi-trasparenti sulla '
         'destra, sotto la parte snella della barra:')
add_bullet('Box GRANDE (con "/" separatore): x=1401-1431 (86.48%-88.33%), centro X=87.40%')
add_bullet('Box PICCOLO (per LV): x=1432-1440 (88.40%-88.89%), centro X=88.65%')
add_bullet('Y entrambi: 68.22%-84.50%, centro Y=76.36%')
add_para('Implementazione valori:', bold=True)
add_bullet('I numeri (300/300, 120/120, 200/200) sono posizionati DENTRO il box GRANDE esistente del PNG')
add_bullet('Il LV (01) è posizionato DENTRO il box PICCOLO esistente (SOLO per la barra Energy)')
add_bullet('HP e MP: SOLO valori, nessun LV')
add_bullet('Energy: valori + LV (SOLO questa barra ha LV)')
add_bullet('NON creare box aggiuntivi — usare SOLO i box esistenti del PNG')
add_bullet('Font size: clamp(0.35rem, 0.65vw, 0.5rem) per stare nei box')

add_para('Nome giocatore:', bold=True)
add_bullet('Posizionato SOPRA la barra HP (solo su quella)')
add_bullet('Stile: semi-trasparente scuro rgba(48,48,48,0.78), corner tagliati, sottile')
add_bullet('Prende quasi tutta la larghezza della barra HP (left: 8, right: 1)')
add_bullet('NON si interseca con la scritta "HP" (che è in alto a sinistra)')

add_para('Etichette HP/MP/EN:', bold=True)
add_bullet('Posizionate SOPRA ogni barra (small, top-left)')
add_bullet('Colorate per matchare la barra (verde/blu/giallo)')
add_bullet('NON si intersecano con il fill colorato')

add_para('Animazioni:', bold=True)
add_bullet('Fill animato via clip-path: inset(0 X% 0 0) — duration 1.1s easeOut')
add_bullet('Pulse glow quando HP < 25% (warning state)')

add_heading('5.2 SaoMainMenu (menu cascata)', level=2)
add_para('Posizione: top-right (fixed). Singola icona Config.svg.')
add_para('Cliccando l\'icona si apre un menu a cascata verticale con 9 voci:')
add_code('''1. Personaggio
2. Borsa
3. Inventario
4. Quest
5. Piano
6. Party
7. Opzioni
8. Messaggi
9. Log Out''')

add_para('Animazione (da MenuView.qml del repo Progetto-SAO):', bold=True)
add_bullet('OPEN: opacity 0→1 (400ms) + y -height→0 (600ms, easing OutQuart)')
add_bullet('Sound Popup.SAO.Menu.wav dopo 300ms delay')
add_bullet('CLOSE: opacity 1→0 (400ms) + y 0→-height (300ms, easing InQuad)')
add_bullet('Staggered cascade: items appaiono uno dopo l\'altro (60ms delay)')

add_para('Stile voci menu:', bold=True)
add_bullet('Background: #FBFBFB (bianco SAO) di default')
add_bullet('Hover: gradiente blu #5CC4F0→#2B73B3 con testo bianco')
add_bullet('Corner tagliati (clip-path 10px)')
add_bullet('Icona circolare a sinistra (solo se item.icon presente)')

add_heading('5.3 SaoNotificationWindow (finestra notifiche)', level=2)
add_para('Finestra modale centrata. Usa i 3 SVG canonici della cartella "Finestra notifiche SAO":')
add_table(
    ['SVG', 'Dimensioni', 'Contenuto'],
    [
        ['Pezzo superiore finestra.svg', '1200x228', 'Header: gradiente #EFEFEF→#DFDFDF + separatore #A8A8A8. Contiene il TITOLO'],
        ['Parte interna finestra.svg', '1200x300', 'Body: #D6D6D6 con ombre top/bottom. Contiene il TESTO del messaggio'],
        ['Parete sotto della finestra.svg', '1200x330', 'Footer: bianco con 2 cerchi. Cerchio blu a sinistra (330,165), cerchio rosso a destra (870,165)'],
    ]
)

add_para('Layout (proporzioni):', bold=True)
add_bullet('TOP: 26.6% (228/858) — header con titolo')
add_bullet('MIDDLE: 35.0% (300/858) — body con testo')
add_bullet('BOTTOM: 38.4% (330/858) — footer con bottoni')
add_bullet('Dimensione finestra: min(520px, 92vw), aspectRatio 1200/858')

add_para('Bottoni (NO scritte sopra, si capiscono da soli):', bold=True)
add_bullet('Bottone BLU (sinistra, x=27.5%): CONFERMA — anello con punto centrale')
add_bullet('Bottone ROSSO (destra, x=72.5%): ANNULLA/CHIUDI — disco pieno con X bianca')
add_bullet('Hover effect: semplice "lift" (y: -4px + scale 1.05 + ombra forte sotto)')
add_bullet('NON usare glow eccessivi o anelli rotanti (rimossi su richiesta utente)')

add_para('Animazione (da MenuView.qml):', bold=True)
add_bullet('OPEN: opacity 0→1 (400ms) + y -height→0 (600ms, OutQuart)')
add_bullet('Sound Popup.SAO.Message.wav dopo 300ms')
add_bullet('CLOSE: opacity 1→0 (400ms) + y 0→-height (300ms, InQuad)')

add_para('Tipi di notifica (kind):', bold=True)
add_table(
    ['Kind', 'Colore', 'Icona', 'Suono'],
    [
        ['system', '#2B73B3 (blu)', '/sao/menu/Config.svg', 'Notify.SAO.System.wav'],
        ['message', '#5CC4F0 (cyan)', '/sao/menu/Message.svg', 'Notify.SAO.Message.wav'],
        ['alert', '#BE2156 (rosso)', '/sao/hex/Warning.svg', 'Notify.SAO.Alert.wav'],
        ['present', '#EBA601 (giallo)', '/sao/hex/SAO_Congratulations!!.svg', 'Notify.SAO.Present.wav'],
    ]
)

add_heading('5.4 ParticleBackground (background 3D)', level=2)
add_para('Background fisso con React Three Fiber.')
add_bullet('1500 particelle in volume 3D con shader custom (glow + twinkle)')
add_bullet('2 point lights che seguono il mouse (cyan #2B73B3 + bianco)')
add_bullet('Repulsione particelle vicino al cursore (raggio 8 unità)')
add_bullet('Parallax rotazione campo (segue mouse)')
add_bullet('Spotlight boost da hover delle card UI')
add_bullet('Lazy DPR per mobile ([1, 1.2] vs [1, 2])')
add_bullet('Volumetric planes per depth')
add_bullet('Fog: #020814, 25-55')

add_heading('5.5 LoginScreen', level=2)
add_para('Boot sequence VR multi-stage con log lines:')
add_code('''
'> INITIALIZING NERVEGEAR v1.100...'
'> CHECKING BRAIN-WAVE INTERFACE... OK'
'> CALIBRATING VR VISUAL CORTEX... OK'
'> SYSTEM READY — WELCOME, PLAYER' ''')
add_bullet('Suoni: Startup.SAO.NerveGear → system → popupPanel → welcome (sfalsati)')
add_bullet('Logo NerveGear: w-[280px] sm:w-[420px] md:w-[520px] con glow pulsante 3s')
add_bullet('Parallax: logo e card si muovono col mouse (x*15, y*15)')
add_bullet('Pulsanti anelli concentrici attorno al NerveGear')

add_heading('5.6 CharacterCreation', level=2)
add_para('SAO è skill-based — NESSUNA CLASCE. Solo genere + nome.')
add_bullet('Step 1: scelta genere (2 card SAO_Man / SAO_Woman)')
add_bullet('Step 2: riepilogo con nome + stats (tutte a 1)')
add_bullet('Card con clip-path angolare (16px), 3D tilt hover')
add_bullet('Statistiche iniziali: STR=1, VIT=1, AGI=1, DEX=1, INT=1, MEN=1, RES=1')
add_bullet('NON ha step classe (SAO non ha classi)')

doc.add_page_break()

# ===================================================================
# 6. ASSET UTILIZZATI
# ===================================================================
add_heading('6. Asset Utilizzati (esclusivamente dai repo GitHub)', level=1)

add_para('TUTTI gli asset provengono da questi 2 repo (NON creare MAI grafiche da zero):')
add_code('Repo 1: https://github.com/emanuelpapasidero98-crypto/asset-gioco-di-SAO.git\n'
         'Repo 2: https://github.com/emanuelpapasidero98-crypto/Progetto-SAO.git')

add_heading('6.1 SVG — asset-gioco-di-SAO', level=2)
add_table(
    ['Cartella', 'File', 'Uso nel progetto'],
    [
        ['Root', 'SAO_UI_v1.100.svg', 'Riferimento UI generale'],
        ['Root', 'SAO_Man.svg', 'Avatar maschile (USA QUESTO)'],
        ['Root', 'SAO_Woman.svg', 'Avatar femminile (USA QUESTO)'],
        ['Root', 'SAO_Nervegear.svg', 'Logo NerveGear (login + entering)'],
        ['SAO_Login_UI_v1.100/', 'SAO_Login_UI_Transparent.svg', 'Pannello login'],
        ['SAO_Login_UI_v1.100/', 'SAO_Login_UI.svg', 'Pannello login (variante)'],
        ['SAO_Login_Tech-Circle_v1.000/', 'SAO_Login_Tech-Circle_v1.000.svg', 'Cerchio tecnologico rotante'],
        ['Pezzi barra HP, Mana e Energia/', '[Blank] 2.png', 'TEMPLATE BARRA CANONICO (1620x258)'],
        ['Pezzi barra HP, Mana e Energia/', 'contenitore Barre.svg', 'Sfondo scuro barra'],
        ['Pezzi barra HP, Mana e Energia/', 'Contenuto barra HP.svg', 'Fill verde'],
        ['Pezzi barra HP, Mana e Energia/', 'Contorno barre.svg', 'Cornice'],
        ['Pezzi barra HP, Mana e Energia/', 'pezzi valori barre e lv.svg', 'Box valori con "/" e "LV:"'],
        ['Finestra notifiche SAO/', 'SAO_UI-Window_blank.svg', 'Finestra completa (NON usare, meglio i 3 pezzi)'],
        ['Finestra notifiche SAO/', 'Pezzo superiore finestra.svg', 'Header finestra'],
        ['Finestra notifiche SAO/', 'Parte interna finestra.svg', 'Body finestra'],
        ['Finestra notifiche SAO/', 'Parete sotto della finestra.svg', 'Footer con bottoni'],
        ['Finestra notifiche SAO/', 'Pulsante azzurro.svg', 'Bottone OK blu'],
        ['Finestra notifiche SAO/', 'Pulsante rosso.SVG', 'Bottone X rosso'],
        ['1_Menu-1/', 'Man.svg + Man_on.svg', 'Icona Personaggio'],
        ['1_Menu-1/', 'Location.svg + Location_on.svg', 'Icona Mappa'],
        ['1_Menu-1/', 'Message.svg + Message_on.svg', 'Icona Messaggi'],
        ['1_Menu-1/', 'Config.svg + Config_on.svg', 'Icona Impostazioni / Menu'],
        ['1_Menu-1/Icone statistiche/', 'Forza.png, Vita.png, ecc.', '7 icone statistiche'],
        ['SAO_HEX/', 'Warning.svg, You_are_dead.svg, ecc.', 'Avvisi esagonali'],
        ['Background e sfondi/', 'Aincrad.png', 'Città volante (game background)'],
        ['Background e sfondi/', 'Città dell\'inizio.png', 'Città iniziale'],
        ['Background e sfondi/', 'Background Ui.png', 'Background UI generico'],
        ['Equipment Icons/', 'icon_sword.png, icon_bow.png, ecc.', '24 icone equipaggiamento'],
        ['Status Icons/', 'icon_status_*.png', '27 icone status effects'],
    ]
)

add_heading('6.2 Audio — Progetto-SAO', level=2)
add_table(
    ['File', 'Uso nel progetto'],
    [
        ['Startup.SAO.NerveGear.wav', 'Boot login (all\'avvio)'],
        ['Ready.SAO.Welcome.wav', 'Boot login (dopo NerveGear)'],
        ['LinkStart.SAO.Kirito.wav', 'Transizione Link Start (dopo login)'],
        ['LinkStart.SAO.Asuna.wav', 'Transizione entering (dopo creation)'],
        ['Feedback.SAO.Click.wav', 'Hover/click generale'],
        ['Popup.SAO.Menu.wav', 'Apertura menu cascata (dopo 300ms)'],
        ['Popup.SAO.Panel.wav', 'Apparizione HUD/pannelli'],
        ['Popup.SAO.Message.wav', 'Apparizione finestra notifiche (dopo 300ms)'],
        ['Popup.SAO.Launcher.wav', 'Apertura launcher'],
        ['Notify.SAO.System.wav', 'Notifica tipo system'],
        ['Notify.SAO.Message.wav', 'Notifica tipo message'],
        ['Notify.SAO.Alert.wav', 'Notifica tipo alert'],
        ['Notify.SAO.Warning.wav', 'Avvisi di pericolo'],
        ['Notify.SAO.Present.wav', 'Notifica tipo present'],
        ['Dismiss.SAO.Launcher.wav', 'Chiusura launcher/menu'],
        ['Dismiss.SAO.Message.wav', 'Chiusura notifica'],
        ['ProgramStart.wav', 'Avvio programma'],
        ['ProgramReady.wav', 'Sistema pronto'],
        ['Credits.wav', 'Crediti'],
    ]
)

add_heading('6.3 Font — Progetto-SAO/Fonts/', level=2)
add_bullet('SAOUI-Regular.otf — il font canonico di SAO (obbligatorio per tutti i testi)')
add_bullet('Registrato in globals.css con @font-face')
add_bullet('Applicato a TUTTO il body nel layout.tsx')
add_bullet('Nome font: \'SAO UI\' (con fallback Trebuchet MS, sans-serif)')

doc.add_page_break()

# ===================================================================
# 7. SISTEMA AUDIO
# ===================================================================
add_heading('7. Sistema Audio (useSaoSound)', level=1)

add_para('Hook personalizzato in src/hooks/useSaoSound.ts')
add_para('18 suoni diversi mappati:', bold=True)
add_code('''type SaoSoundName =
  | 'startup'        | 'click'        | 'linkStartK'  | 'linkStartA'
  | 'welcome'        | 'alert'        | 'system'      | 'warning'
  | 'present'        | 'message'      | 'popupMenu'   | 'popupPanel'
  | 'popupMessage'   | 'popupLauncher'| 'dismissLauncher' | 'dismissMessage'
  | 'programStart'   | 'programReady' | 'credits';''')

add_para('Caratteristiche:', bold=True)
add_bullet('Cache audio con Map<string, HTMLAudioElement>')
add_bullet('Clone per playback sovrapposto (es. hover veloci)')
add_bullet('Gestione autoplay bloccato (silent fail)')
add_bullet('Preload di tutti i suoni on mount')
add_bullet('API: play(name, volume) + setMuted(boolean)')

doc.add_page_break()

# ===================================================================
# 8. ANIMAZIONI
# ===================================================================
add_heading('8. Animazioni e Interazioni', level=1)

add_heading('8.1 Animazioni canoniche SAO (da MenuView.qml)', level=2)
add_para('Tutte le animazioni di menu/finestre seguono il pattern del file '
         'Progetto-SAO/qml/MenuView.qml (aniFadeIn / aniFadeOut):')
add_code('''OPEN (ParallelAnimation):
  - opacity: 0 → 1 over 400ms
  - y: -height → 0 over 600ms with easing OutQuart (cubic-bezier 0.22,1,0.36,1)
  - play sound after 300ms delay

CLOSE (SequentialAnimation → ParallelAnimation):
  - opacity: 1 → 0 over 400ms
  - y: 0 → -height over 300ms with easing InQuad
  - then visible = false''')

add_heading('8.2 Hover effects', level=2)
add_table(
    ['Elemento', 'Hover effect'],
    [
        ['Card genere (creation)', '3D tilt (rotateX/Y) + spotlight radiale che segue cursore + scale icone'],
        ['Menu items (cascata)', 'Scale 1.05 + background gradiente blu + glow'],
        ['Bottoni blu/rosso (notifica)', 'Semplice "lift" (y: -4px + scale 1.05 + ombra forte) — NON glow eccessivi'],
        ['Icona menu (top-right)', 'Scale 1.1 + glow ring + rotating dashed ring quando aperto'],
        ['Barre HP/MP/Energia', 'Pulse glow quando < 25% (warning state)'],
    ]
)

add_heading('8.3 Transizioni stage', level=2)
add_bullet('Fade in/out duration 0.6s tra stage')
add_bullet('AnimatePresence mode="wait" per evitare overlap')
add_bullet('Link Start: anelli concentrici espandenti + flash bianco finale')

add_heading('8.4 Background particellare', level=2)
add_bullet('1500 particelle con shader custom (vertex + fragment)')
add_bullet('Twinkle: 0.7 + 0.3 * sin(time * 1.5 + position)')
add_bullet('Repulsione mouse: raggio 8 unità, force 1.5x')
add_bullet('2 point lights che seguono il mouse (intensity 6 + 2.5)')
add_bullet('Spotlight boost da UI hover (intensity += 2)')

doc.add_page_break()

# ===================================================================
# 9. FONT SAO UI
# ===================================================================
add_heading('9. Font SAO UI', level=1)

add_para('Il font canonico SAO è SAOUI-Regular.otf, trovato in:')
add_code('Progetto-SAO/Fonts/SAOUI-Regular.otf')

add_para('Registrato in src/app/globals.css:', bold=True)
add_code('''@font-face {
  font-family: 'SAO UI';
  src: url('/sao/fonts/SAOUI-Regular.otf') format('opentype');
  font-weight: normal;
  font-style: normal;
  font-display: swap;
}''')

add_para('Applicato globalmente nel layout.tsx:', bold=True)
add_code('''<body style={{ fontFamily: "'SAO UI', 'Trebuchet MS', sans-serif" }}>''')

add_para('Regole per i testi:', bold=True)
add_bullet('TUTTI i testi usano font SAO UI')
add_bullet('Font weight: 400 (normale) — NON grassetto eccessivo')
add_bullet('Letter-spacing: 0.2em-0.45em per titoli (tracking ampio stile SAO)')
add_bullet('Colori testo: #FBFBFB (bianco) su sfondo scuro, #1a2a3a (scuro) su sfondo chiaro')

doc.add_page_break()

# ===================================================================
# 10. PALETTE COLORI
# ===================================================================
add_heading('10. Palette Colori SAO', level=1)

add_table(
    ['Colore', 'Hex', 'Uso'],
    [
        ['Blu principale', '#0682BE', 'Pannello login (fill)'],
        ['Blu bottoni/cerchi', '#2B73B3', 'Bottoni, bordi, accenti'],
        ['Cyan', '#5CC4F0', 'Glow, icone hover, testo secondario'],
        ['Bianco SAO', '#FBFBFB', 'Sfondi finestre, testo su scuro'],
        ['Grigio chiaro', '#EFEFEF / #DFDFDF', 'Header finestre (gradiente)'],
        ['Grigio medio', '#D6D6D6', 'Body finestre'],
        ['Grigio separatore', '#A8A8A8', 'Separatori finestre'],
        ['Grigio bordo', '#373737', 'Bordo barre'],
        ['Grigio scuro', '#303030', 'Box valori (background)'],
        ['Nero', '#151515', 'Bordi box valori'],
        ['Metallico', '#7A7A7A / #5a5a5a', 'Bordi metallici'],
        ['Rosso SAO', '#BE2156', 'Bottone X, alerts'],
        ['Giallo skill', '#EBA601', 'LV badge, highlights'],
        ['Verde HP', '#7FC522', 'Barra HP (con riflesso #AEDB5E)'],
        ['Scuro testo', '#1a2a3a', 'Testo su sfondo chiaro'],
    ]
)

doc.add_page_break()

# ===================================================================
# 11. DECISIONI DI DESIGN
# ===================================================================
add_heading('11. Decisioni di Design Prese', level=1)

decisions = [
    ('Avatar', 'Usare SAO_Man.svg/SAO_Woman.svg (NON Avatar_maschio.svg/femmina.svg — non piaciuti)'),
    ('Classi', 'RIMOSSE — SAO è skill-based, nessuna classe. Solo genere + nome'),
    ('Statistiche', 'Tutte a 1 (non 8) — baseline minimo, crescita tramite gioco'),
    ('Barre HP/MP/Energia', 'Usare [Blank] 2.png come template. Valori DENTRO i box esistenti del PNG (NON creare box)'),
    ('LV', 'SOLO nella barra Energy (box piccolo a destra). HP/MP non hanno LV'),
    ('Font', 'SAO UI (SAOUI-Regular.otf) ovunque. Weight 400, non grassetto eccessivo'),
    ('Scritta NERVEGEAR', 'RIMOSSA da login e entering (era troppo preponderante)'),
    ('"Il mondo di spada e magia"', 'RIMOSSO da entering'),
    ('"CREAZIONE PERSONAGGIO / FLOOR 1"', 'RIMOSSO header da character creation'),
    ('Tooltip MENU', 'RIMOSSO da SaoMainMenu (l\'icona si capisce da sola)'),
    ('Skill quick-slot', 'RIMOSSO da GameScreen (non richiesto)'),
    ('Demo HP drain', 'RIMOSSO — le barre non devono diminuire da sole'),
    ('Bottoni notifica', 'NO scritte "OK"/"SÌ"/"NO" — i cerchi si capiscono da soli'),
    ('Animazione hover bottoni', 'Semplice "lift" (NON glow/anello rotante)'),
    ('Finestra notifica', 'Usare i 3 SVG separati (NON SAO_UI-Window_blank.svg intero)'),
    ('Nome giocatore', 'Stile semi-trasparente scuro (NON bianco SAO)'),
    ('Card gender', 'Corner tagliati (clip-path 16px) stile SAO, NON arrotondate'),
    ('Background login', 'Particle field 3D + scanlines + vignette + chromatic aberration'),
]

for argomento, decisione in decisions:
    p = doc.add_paragraph()
    run = p.add_run(f'{argomento}: ')
    run.bold = True
    run.font.color.rgb = SAO_BLUE
    p.add_run(decisione)

doc.add_page_break()

# ===================================================================
# 12. REGOLE FONDAMENTALI
# ===================================================================
add_heading('12. Regole Fondamentali (NON rompere MAI)', level=1)

add_para('Queste regole sono SACRE. Non devono MAI essere violate:', bold=True, color=RGBColor(0xBE, 0x21, 0x56))

rules = [
    'NON CREARE MAI GRAFICHE DA ZERO. Usa SOLO asset dai repo GitHub.',
    'Usa SEMPRE il font SAO UI (SAOUI-Regular.otf) per tutti i testi.',
    'Usa SEMPRE i colori canonici SAO (#2B73B3, #FBFBFB, #EBA601, #BE2156, #7FC522).',
    'Le barre HP/MP/Energia usano il PNG [Blank] 2.png come template.',
    'I valori delle barre vanno DENTRO i box esistenti del PNG (NON creare box).',
    'Il LV va SOLO nella barra Energy (nel box piccolo a destra).',
    'SAO NON ha classi — è skill-based. Non reintroitare classi.',
    'Le animazioni di menu/finestre seguono il pattern di MenuView.qml (fade + slide dall\'alto).',
    'La finestra notifica usa i 3 SVG canonici (top + middle + bottom), NON ricrearla.',
    'I bottoni blu/rosso della finestra non hanno scritte (si capiscono da soli).',
    'L\'hover dei bottoni è un semplice "lift" (NON glow eccessivi o anelli rotanti).',
    'Il menu in alto a destra è UNA sola icona che apre una cascata di 9 voci.',
    'Tutti i testi sono in italiano.',
    'I font weight sono 400 (non grassetto eccessivo).',
    'Le card SAO hanno corner tagliati (clip-path), NON arrotondati.',
]
for rule in rules:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run('⚠ ' + rule)
    run.font.color.rgb = RGBColor(0xBE, 0x21, 0x56)
    run.bold = True

doc.add_page_break()

# ===================================================================
# 13. PROBLEMI RISOLTI
# ===================================================================
add_heading('13. Problemi Risolti e Soluzioni', level=1)

problems = [
    ('Testo login invisibile (bianco su bianco)', 'Cambiato colore testo input da bianco a #1a2a3a (scuro) per contrastare il sfondo bianco del SVG'),
    ('Barre HP non centravano i valori', 'Analizzato il PNG con Python per trovare le coordinate esatte dei box semi-trasparenti (x=86.48%-88.33% per box grande, x=88.40%-88.89% per box LV)'),
    ('Tasto rosso non funzionava', 'Aggiunto handler onCancel separato (prima era condizionato a cancelLabel)'),
    ('Finestra notifica disallineata', 'Sostituito il singolo SVG SAO_UI-Window_blank.svg con i 3 pezzi separati sovrapposti'),
    ('Avatar troppo grandi/ brutti', 'Tornati a SAO_Man.svg/SAO_Woman.svg (Avatar_maschio/femmina.svg non piaciuti)'),
    ('Glow eccessivo sui bottoni', 'Sostituito con semplice "lift" (y: -4px + scale + ombra)'),
    ('NerveGear troppo preponderante', 'Rimossa la scritta NERVEGEAR bianca gigante (tenuto solo il logo)'),
    ('Skill panel non richiesto', 'Rimosso da GameScreen'),
    ('Demo HP drain automatico', 'Rimosso — le barre cambiano solo via user actions'),
    ('Scritta MENU tooltip', 'Rimossa — l\'icona si capisce da sola'),
    ('Background troppo scuro', 'Aggiunto Aincrad.png come background con overlay per leggibilità'),
    ('Font grassetto eccessivo', 'Cambiato tutto a fontWeight: 400'),
]

for problema, soluzione in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'Problema: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xBE, 0x21, 0x56)
    p.add_run(problema)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'Soluzione: ')
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x2B, 0x73, 0xB3)
    p2.add_run(soluzione)
    doc.add_paragraph()

doc.add_page_break()

# ===================================================================
# 14. ROADMAP
# ===================================================================
add_heading('14. Roadmap — Prossimi Sviluppi', level=1)

add_para('Feature da implementare prossimamente (in ordine di priorità):')

roadmap = [
    ('ALTA', 'Sistema equipaggiamento', 'Usare gli Equipment Icons (spada, pugnale, arco, due lame, ecc.) per creare un inventario equipaggiamento'),
    ('ALTA', 'Inventario completo', 'Usare gli Item Types icons per un inventario a griglia'),
    ('ALTA', 'Status effects', 'Usare le 27 Status Icons per effetti di stato (veleno, stun, burn, ecc.)'),
    ('ALTA', 'Combat system a turni', 'Sistema di combattimento RPG a turni con HP/MP/Energy'),
    ('MEDIA', 'Mappa Aincrad', 'Sistema di navigazione tra i piani con Location.svg'),
    ('MEDIA', 'Sistema skill', 'Skill tree sbloccabile (SAO è skill-based, no classi)'),
    ('MEDIA', 'Party system', 'Formazione party con altri giocatori (NPC per ora)'),
    ('MEDIA', 'Quest system', 'Log delle quest con tracker'),
    ('MEDIA', 'Borsa/wallet', 'Sistema valuta (Col) e transazioni'),
    ('BASSA', 'Messaggi', 'Sistema di messaggi tra giocatori'),
    ('BASSA', 'Opzioni', 'Menu impostazioni (audio, grafica, controlli)'),
    ('BASSA', 'Salvataggio', 'Persistenza dati con Prisma/SQLite'),
]

add_table(
    ['Priorità', 'Feature', 'Descrizione'],
    [[p, f, d] for p, f, d in roadmap]
)

doc.add_page_break()

# ===================================================================
# 15. COME CLONARE
# ===================================================================
add_heading('15. Come Clonare e Avviare il Progetto', level=1)

add_para('Repository GitHub:', bold=True)
add_code('https://github.com/emanuelpapasidero98-crypto/Project-SAO-Alpha.git')

add_para('Istruzioni:', bold=True)
add_code('''# 1. Clona il repository
git clone https://github.com/emanuelpapasidero98-crypto/Project-SAO-Alpha.git
cd Project-SAO-Alpha

# 2. Installa le dipendenze
bun install

# 3. Avvia il dev server
bun run dev

# 4. Apri http://localhost:3000 nel browser''')

add_para('Note:', bold=True)
add_bullet('Richiede Node.js 18+ e Bun installati')
add_bullet('Il dev server si avvia automaticamente sulla porta 3000')
add_bullet('Tutti gli asset sono inclusi nel repo (in public/sao/)')
add_bullet('Il font SAO UI è in public/sao/fonts/SAOUI-Regular.otf')

doc.add_page_break()

# ===================================================================
# 16. ISTRUZIONI PER CLAUDE
# ===================================================================
add_heading('16. Istruzioni per Claude (AI assistant)', level=1)

add_para('Se stai leggendo questo documento per continuare lo sviluppo del progetto SAO Alpha, ecco le linee guida fondamentali:', bold=True)

add_heading('Context del progetto', level=2)
add_bullet('SAO Alpha è una web app RPG che riproduce fedelmente l\'UI di Sword Art Online')
add_bullet('Il progetto è in stato Alpha — login + character creation + game screen HUD sono completi')
add_bullet('Il prossimo step è probabilmente il combat system o l\'inventario')
add_bullet('L\'utente (Emanuele) è molto esigente sull\'estetica canonica SAO')

add_heading('Regole d\'oro per continuare lo sviluppo', level=2)
add_bullet('NON CREARE MAI GRAFICHE DA ZERO — usa solo asset dai repo GitHub elencati')
add_bullet('Prima di aggiungere una feature, controlla SEMPRE se esiste già un asset nel repo')
add_bullet('Le animazioni devono seguire il pattern di MenuView.qml (fade + slide dall\'alto)')
add_bullet('I colori devono essere della palette SAO canonica')
add_bullet('Il font deve essere SEMPRE SAO UI')
add_bullet('Le card/finestre devono avere corner tagliati (clip-path), non arrotondati')
add_bullet('I testi sono in italiano')
add_bullet('Non usare grassetto eccessivo (fontWeight: 400)')

add_heading('Asset repo URLs', level=2)
add_code('Repo 1: https://github.com/emanuelpapasidero98-crypto/asset-gioco-di-SAO.git\n'
         'Repo 2: https://github.com/emanuelpapasidero98-crypto/Progetto-SAO.git\n'
         'Project: https://github.com/emanuelpapasidero98-crypto/Project-SAO-Alpha.git')

add_heading('File chiave da leggere prima di modificare', level=2)
add_code('''src/app/page.tsx                    # State machine principale
src/components/sao/SaoHUD.tsx       # Barre HP/MP/Energia (complesso)
src/components/sao/GameScreen.tsx   # Schermata di gioco
src/components/sao/SaoMainMenu.tsx  # Menu cascata
src/components/sao/SaoNotificationWindow.tsx # Finestra notifiche
src/hooks/useSaoSound.ts            # Audio manager
src/lib/sao-data.ts                 # Dati personaggio (stats, gender)
src/app/globals.css                 # Font SAO UI @font-face''')

add_heading('Cose da NON fare', level=2, color=RGBColor(0xBE, 0x21, 0x56))
add_bullet('NON reintroitare classi personaggio (SAO è skill-based)')
add_bullet('NON usare Avatar_maschio.svg/Avatar_femmina.svg (usare SAO_Man/Woman.svg)')
add_bullet('NON aggiungere badge LV sotto le barre (solo nel box Energy)')
add_bullet('NON creare box valori separati per le barre (usare quelli del PNG)')
add_bullet('NON usare glow eccessivi o anelli rotanti per hover bottoni (usare "lift")')
add_bullet('NON aggiungere scritte sui bottoni blu/rosso della finestra notifica')
add_bullet('NON far diminuire le barre automaticamente (solo via user actions)')

add_heading('Come testare le modifiche', level=2)
add_bullet('Usa Agent Browser per navigare: agent-browser open http://localhost:3000')
add_bullet('Usa VLM (z-ai vision) per verificare l\'aspetto visivo')
add_bullet('Controlla sempre la console per errori: agent-browser console')
add_bullet('Verifica il flusso completo: login → linkstart → creation → entering → game')

doc.add_page_break()

# ===================================================================
# APPENDICE
# ===================================================================
add_heading('Appendice A — Coordinate slot barre HP', level=1)
add_para('Coordinate esatte dei box semi-trasparenti nel PNG [Blank] 2.png (1620x258):')
add_table(
    ['Elemento', 'X range', 'Y range', 'Centro', 'Note'],
    [
        ['Barra (parte larga)', '21%-93%', '23.6%-38%', '—', 'Fill colorato principale'],
        ['Barra (parte snella)', '21%-58%', '38%-51.2%', '—', 'Parte inferiore più stretta'],
        ['Box valori GRANDE', '86.48%-88.33%', '68.22%-84.50%', 'X=87.40%, Y=76.36%', 'Contiene "/" e i valori 300/300'],
        ['Box LV PICCOLO', '88.40%-88.89%', '68.22%-84.50%', 'X=88.65%, Y=76.36%', 'Contiene LV (solo per Energy)'],
        ['Separatore "/"', '87.78%-88.02%', '68.22%-84.50%', '—', 'Linea verticale nel box grande'],
    ]
)

add_heading('Appendice B — Worklog completo', level=1)
add_para('Il worklog completo di tutte le sessioni di sviluppo è nel file worklog.md nel repository.')

# Save
doc.save(OUTPUT)
import os
size = os.path.getsize(OUTPUT)
print(f'Documento generato: {OUTPUT}')
print(f'Dimensione: {size/1024:.1f} KB')
